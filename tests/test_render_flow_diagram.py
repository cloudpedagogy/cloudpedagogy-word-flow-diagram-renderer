import tempfile
import unittest
from pathlib import Path

from docx import Document

import render_flow_diagram as flow


def make_doc(path: Path, nodes, links, node_headers=None, link_headers=None):
    doc = Document()
    doc.add_heading("NODES", 1)
    table = doc.add_table(rows=1, cols=len(node_headers or ["ID", "Label", "Type"]))
    for cell, value in zip(table.rows[0].cells, node_headers or ["ID", "Label", "Type"]):
        cell.text = value
    for values in nodes:
        for cell, value in zip(table.add_row().cells, values):
            cell.text = value
    doc.add_heading("LINKS", 1)
    table = doc.add_table(rows=1, cols=len(link_headers or ["Source", "Target", "Label"]))
    for cell, value in zip(table.rows[0].cells, link_headers or ["Source", "Target", "Label"]):
        cell.text = value
    for values in links:
        for cell, value in zip(table.add_row().cells, values):
            cell.text = value
    doc.save(path)


class FlowTests(unittest.TestCase):
    def parse(self, nodes, links, node_headers=None, link_headers=None):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "input.docx"
            make_doc(path, nodes, links, node_headers, link_headers)
            return flow.read_docx(path)

    def test_valid_flow_and_aliases(self):
        result = self.parse(
            [["a", "Start", "Start"], ["b", "Review", "Decision"]],
            [["a", "b", "Continue"]],
            ["Key", "Name", "Kind"], ["From", "To", "Condition"],
        )
        self.assertEqual(2, len(result.nodes))
        self.assertEqual("decision", result.nodes[1]["type"])
        self.assertEqual(1, len(result.links))
        self.assertFalse([i for i in result.issues if i.level == "error"])

    def test_duplicate_node_removed(self):
        result = self.parse([["a", "A", "Process"], ["a", "Again", "Process"]], [])
        self.assertEqual(1, len(result.nodes))
        self.assertTrue(any("Duplicate node" in i.message for i in result.issues))

    def test_missing_link_node_removed(self):
        result = self.parse([["a", "A", "Process"]], [["a", "missing", "Next"]])
        self.assertFalse(result.links)
        self.assertTrue(any("missing node" in i.message for i in result.issues))

    def test_unsafe_link_and_bad_colour_cleaned(self):
        result = self.parse(
            [["a", "A", "Process", "javascript:alert(1)", "#bad-color"]],
            [],
            ["ID", "Label", "Type", "Link", "Colour"],
        )
        self.assertEqual("", result.nodes[0]["link"])
        self.assertEqual(flow.DEFAULT_COLORS["process"], result.nodes[0]["color"])

    def test_unknown_type_becomes_process(self):
        result = self.parse([["a", "A", "Mystery"]], [])
        self.assertEqual("process", result.nodes[0]["type"])
        self.assertTrue(any("Unknown node type" in i.message for i in result.issues))

    def test_mermaid_escapes_labels(self):
        result = self.parse([["a", 'A "quoted" <node>', "Process"]], [])
        diagram, _ = flow.diagram_text(result, "LR")
        self.assertIn("&quot;quoted&quot;", diagram)
        self.assertIn("&lt;node&gt;", diagram)


if __name__ == "__main__":
    unittest.main()
