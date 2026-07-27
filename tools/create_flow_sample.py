#!/usr/bin/env python3
"""Create the populated Word input used by the flow diagram demonstration."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).parents[1] / "examples/flow_diagram_example.docx"
BLUE, DARK, PALE = "2E74B5", "1F4D78", "E8EEF5"


def set_font(run, size=9, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size, run.bold, run.font.color.rgb = Pt(size), bold, RGBColor.from_string(color)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def set_cell_margins(cell):
    props = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    props.append(margins)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        shade(cell, PALE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            set_font(run, 8.5, True, DARK)
    header_props = table.rows[0]._tr.get_or_add_trPr()
    header_props.append(OxmlElement("w:tblHeader"))
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_font(run, 8.25)
        table.rows[-1]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    set_table_geometry(table, widths)
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(.65)
    sec.left_margin = sec.right_margin = Inches(.65)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(10)
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.1
    for name, size, before, after in (("Title", 22, 0, 6), ("Heading 1", 15, 12, 6)):
        style = doc.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(size)
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)

    doc.add_paragraph("Word-to-Interactive Flow Diagram: Sample Input", style="Title")
    doc.add_paragraph(
        "Edit these tables, save the document, then run render_flow_diagram.py. "
        "Node IDs must be unique; every Source and Target value must match a node ID."
    )
    doc.add_heading("SETTINGS", level=1)
    add_table(doc, ["Setting", "Value"], [
        ["title", "Online Course Development Workflow"],
        ["subtitle", "A generic quality-assured route from proposal to Moodle publication"],
        ["orientation", "LR"],
        ["theme", "light"],
        ["show_legend", "true"],
        ["allow_orientation_switching", "true"],
    ], [2600, 11100])

    doc.add_heading("NODES", level=1)
    nodes = [
        ["start", "Course proposal", "Start", "Initial course or module proposal is agreed.", "Approved", "#DCFCE7", "", "Planning"],
        ["draft", "Academic drafting", "Process", "Academic author prepares structured learning content.", "In progress", "#DBEAFE", "", "Development"],
        ["review", "Digital education review", "Process", "Learning design, structure and interaction review.", "Pending", "#DBEAFE", "", "Development"],
        ["revise", "Revisions required?", "Decision", "Decide whether substantive changes are needed.", "Decision", "#FEF3C7", "", "Review"],
        ["update", "Revise content", "Process", "Author responds to review findings.", "Pending", "#FEE2E2", "", "Development"],
        ["access", "Accessibility check", "Process", "Check headings, links, tables, media and alternatives.", "Pending", "#EDE9FE", "", "Assurance"],
        ["approve", "Academic approval", "Decision", "Confirm that content is accurate and ready to release.", "Pending", "#FEF3C7", "", "Assurance"],
        ["publish", "Publish to Moodle", "Process", "Create the learner-facing version and verify settings.", "Pending", "#DBEAFE", "https://moodle.org/", "Publication"],
        ["qa", "Post-publication QA", "Document", "Record the final quality assurance outcome.", "Pending", "#FCE7F3", "", "Publication"],
        ["end", "Course available", "End", "The approved course is available to learners.", "Ready", "#DCFCE7", "", "Publication"],
    ]
    add_table(doc, ["ID", "Label", "Type", "Description", "Status", "Colour", "Link", "Group"], nodes,
              [950, 1900, 1050, 3900, 1100, 1050, 1800, 1950])

    doc.add_heading("LINKS", level=1)
    links = [
        ["start", "draft", "Begin", "solid", "", "Start content production."],
        ["draft", "review", "Submit", "solid", "", "Send draft for review."],
        ["review", "revise", "", "solid", "", "Review produces a decision."],
        ["revise", "update", "Yes", "dashed", "#DC2626", "Changes are required."],
        ["update", "review", "Resubmit", "dashed", "#DC2626", "Return revised content for review."],
        ["revise", "access", "No", "solid", "#15803D", "Proceed to accessibility checking."],
        ["access", "approve", "Pass", "solid", "", "Send accessible version for approval."],
        ["approve", "update", "Not yet", "dashed", "#DC2626", "Return to revision."],
        ["approve", "publish", "Approved", "thick", "#15803D", "Proceed to publication."],
        ["publish", "qa", "Verify", "solid", "", "Perform checks in the live course."],
        ["qa", "end", "Release", "thick", "#15803D", "Make the course available."],
    ]
    add_table(doc, ["Source", "Target", "Label", "Style", "Colour", "Description"], links,
              [1500, 1500, 1700, 1200, 1300, 7050])
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
